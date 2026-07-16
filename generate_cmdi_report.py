#!/usr/bin/env python3
"""命令注入漏洞修复报告 - 新手入门版"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for level in range(1,4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'; hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, hc="2B579A"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), hc); c._tc.get_or_add_tcPr().append(shd)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
            if ri%2==0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "F0F4F8"); c._tc.get_or_add_tcPr().append(shd)
    return t

# 封面
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('命令注入漏洞修复报告'); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor(0x0A,0x1A,0x3A)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('NEXUS 用户管理系统 · 新手入门版'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━'*40).font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
for k,v in [('项目名称','NEXUS 用户管理系统'),('报告类型','命令注入漏洞修复（新手入门版）'),('仓库地址','https://github.com/jgymh/user-management-system')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{k}:  '); r.bold = True; r.font.size = Pt(11)
    p.add_run(v).font.size = Pt(11)
doc.add_page_break()

# 1. 什么是命令注入
doc.add_heading('1. 什么是命令注入漏洞', level=1)
doc.add_paragraph('命令注入（Command Injection）就是攻击者通过在输入框中"夹带"系统命令，让服务器执行不该执行的命令。')

p = doc.add_paragraph()
r = p.add_run('用生活中的例子理解：')
r.bold = True
doc.add_paragraph('你去网吧，跟网管说"帮我开一下3号机"（正常请求）。但你如果说"帮我开一下3号机，然后把收银台的钱都给我"（命令注入），网管真的照做了——这就是命令注入。')

doc.add_paragraph()
doc.add_paragraph('在我们的 Ping 功能里：')
p = doc.add_paragraph()
r = p.add_run('正常输入：')
r.bold = True
p.add_run('8.8.8.8 → 执行 ping -c 3 8.8.8.8 ✅')
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('恶意输入：')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)
p.add_run('8.8.8.8;id → 执行 ping -c 3 8.8.8.8;id 🔥')
doc.add_paragraph('分号 ; 后面的 id 命令也会被执行，泄露当前用户信息！')

doc.add_page_break()

# 2. 项目中发现的漏洞
doc.add_heading('2. 项目中发现了什么漏洞', level=1)

doc.add_heading('漏洞：/ping 接口命令注入', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🔴 严重')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

doc.add_paragraph('/ping 功能使用 f-string 拼接命令 + shell=True 执行，用户输入直接拼进系统命令。')

p = doc.add_paragraph()
r = p.add_run('攻击方法 1：执行任意命令')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('输入: 8.8.8.8;id')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 执行 ping -c 3 8.8.8.8;id，显示当前用户 uid=0(root)')

p = doc.add_paragraph()
r = p.add_run('攻击方法 2：管道符执行')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('输入: 8.8.8.8|whoami')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 执行 ping -c 3 8.8.8.8|whoami，输出 root')

p = doc.add_paragraph()
r = p.add_run('攻击方法 3：链式执行')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('输入: 8.8.8.8&&ls')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 列出网站根目录所有文件')

doc.add_paragraph()
doc.add_paragraph('修复前代码：')
p = doc.add_paragraph()
r = p.add_run('''# ❌ 漏洞代码
cmd = f"ping -c 3 {ip}"           # f-string拼接
result = subprocess.check_output(
    cmd, shell=True, ...)           # shell=True执行''')
r.font.size = Pt(9); r.font.name = 'Courier New'

doc.add_paragraph('修复后代码：')
p = doc.add_paragraph()
r = p.add_run('''# ✅ 安全代码
import shlex
safe_ip = shlex.quote(ip)           # 转义特殊字符
cmd = ["ping", "-c", "3", safe_ip]  # 参数列表模式
result = subprocess.check_output(
    cmd, ...)                        # shell=False(默认)''')
r.font.size = Pt(9); r.font.name = 'Courier New'; r.font.color.rgb = RGBColor(0x00,0x80,0x00)

doc.add_page_break()

# 3. 漏洞汇总表
doc.add_heading('3. 漏洞汇总表', level=1)
add_table(doc, ['编号','漏洞名称','风险','位置','攻击方式','后果'],
    [
        ['SEC-CMD-01','分号注入','🔴 严重','/ping',';id','执行任意命令'],
        ['SEC-CMD-02','管道符注入','🔴 严重','/ping','\|whoami','执行任意命令'],
        ['SEC-CMD-03','逻辑与注入','🔴 严重','/ping','&&ls','执行任意命令'],
        ['SEC-CMD-04','反引号注入','🔴 严重','/ping','`id`','执行任意命令'],
    ])

doc.add_page_break()

# 4. 修复原理
doc.add_heading('4. 修复原理', level=1)

doc.add_heading('修复 1：去掉 shell=True', level=2)
doc.add_paragraph('shell=True 等于把字符串直接丢给系统 shell 解释执行。去掉 shell=True 后，命令以参数列表形式传递，特殊字符不会被 shell 解释。')

add_table(doc, ['对比项','shell=True','shell=False'],
    [['执行方式','字符串丢给/bin/sh','参数列表传递给系统调用'],
     [';id 解释','分号被当作命令分隔符','分号被当作普通参数传给ping'],
     ['\|whoami','管道符创建子进程','管道符被当作普通参数'],
     ['安全性','❌ 危险','✅ 安全']])

doc.add_heading('修复 2：shlex.quote()', level=2)
doc.add_paragraph('即使用了 shell=True，shlex.quote() 也可以转义特殊字符。但最安全的做法是：去掉 shell=True + 参数列表模式。')

doc.add_heading('修复 3：参数列表模式', level=2)
doc.add_paragraph('''用列表传参，subprocess 会直接调用 execvp 系统函数，不经过 shell 解释：
["ping", "-c", "3", "8.8.8.8;id"]
→ ping 收到的参数是 "8.8.8.8;id"，分号不会被当作命令分隔符''')

doc.add_page_break()

# 5. 修复前后对比
doc.add_heading('5. 修复前后对比', level=1)
add_table(doc, ['检测项','修复前','修复后'],
    [
        ['正常Ping 8.8.8.8','✅ 正常','✅ 正常'],
        [';id 注入','✅ uid=0(root) ❌','❌ ping报错，注入失效 ✅'],
        ['\|whoami 注入','✅ root ❌','❌ ping报错，注入失效 ✅'],
        ['&&ls 注入','✅ 列目录 ❌','❌ ping报错，注入失效 ✅'],
        ['shell=True','是','否（参数列表）'],
        ['shlex.quote()','否','是'],
    ], hc="1A6B3C")

doc.add_page_break()

# 6. 新手总结
doc.add_heading('6. 新手总结', level=1)
tips = [
    ('永远不要用 shell=True', '参数列表模式 ["cmd","arg1","arg2"] 永远比 f"cmd {arg}" 安全'),
    ('shlex.quote() 是第二道防线', '如果必须用 shell=True，先 quote 处理用户输入'),
    ('特殊字符是攻击者的武器', '; | && ` $() 都是 shell 特殊字符，不能让它们经过 shell 解释'),
    ('修复前先验证漏洞', '提交 ;id 如果返回 uid=，说明存在命令注入，修好了再测一次应该报错'),
    ('记住一条原则', '用户输入永远不能直接拼进系统命令！'),
]
for title, desc in tips:
    p = doc.add_paragraph()
    r = p.add_run('✅ ' + title); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph(desc + '\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('记住：shell=True + f-string = 命令注入！参数列表模式 = 安全！')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

output_path = '/home/user/projects/user-mgr/命令注入漏洞修复报告_新手版.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
