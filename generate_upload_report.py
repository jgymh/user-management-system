#!/usr/bin/env python3
"""Upload 安全修复报告 - 新手入门版"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def set_cell_shading(cell, color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shd)

def add_table(doc, headers, rows, header_color="2B579A"):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(10)
        set_cell_shading(cell, header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(9.5)
            if ri % 2 == 0: set_cell_shading(cell, "F0F4F8")
    return table

# ═════ 封面 ═════
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('上传功能安全修复报告'); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor(0x0A,0x1A,0x3A)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('NEXUS 用户管理系统 · 新手入门版'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━' * 40).font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
for k, v in [('项目名称','NEXUS 用户管理系统'),('报告类型','上传漏洞安全修复（新手入门版）'),('仓库地址','https://github.com/jgymh/user-management-system')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{k}:  '); r.bold = True; r.font.size = Pt(11)
    p.add_run(v).font.size = Pt(11)
doc.add_page_break()

# ═════ 1. 上传功能有什么漏洞 ═════
doc.add_heading('1. 上传功能有什么漏洞', level=1)
doc.add_paragraph('文件上传是网站最常见的功能之一，但如果没做好安全检查，就会变成最危险的攻击入口。')
doc.add_paragraph('我们的 NEXUS 系统在上传头像功能里，一开始故意没做任何安全防护。来看看有哪些问题：')

doc.add_heading('漏洞 1：没有检查文件类型', level=2)
doc.add_paragraph('任何文件都能上传——图片、exe 程序、PHP 木马、html 钓鱼页面……')
doc.add_paragraph('攻击者上传一个 "shell.php"，如果服务器执行了它，网站就被控制了。')

doc.add_heading('漏洞 2：文件名没有过滤（路径穿越）', level=2)
p = doc.add_paragraph()
r = p.add_run('攻击者把文件命名为 "../../shell.php"')
r.bold = True
doc.add_paragraph('把文件保存到网站根目录以外的地方，覆盖系统文件，或者保存到可执行目录里。')

doc.add_heading('漏洞 3：同名文件覆盖', level=2)
doc.add_paragraph('两个用户都上传 "avatar.jpg"，后上传的会把前面那个覆盖掉。')

doc.add_heading('漏洞 4：文件名包含特殊字符', level=2)
doc.add_paragraph('中文、空格、特殊符号可能导致保存失败、路径解析错误、甚至破坏服务器。')

doc.add_page_break()

# ═════ 2. 漏洞汇总表 ═════
doc.add_heading('2. 漏洞汇总表', level=1)
add_table(doc, ['编号','漏洞名称','风险等级','攻击举例','后果'],
    [
        ['SEC-UPLOAD-01','文件类型未检查','🔴 严重','上传 PHP webshell','服务器被接管'],
        ['SEC-UPLOAD-02','路径穿越','🔴 严重','文件名 "../../evil.exe"','覆盖系统文件'],
        ['SEC-UPLOAD-03','同名覆盖','🟡 中危','两个 avatar.jpg','别人的头像丢了'],
        ['SEC-UPLOAD-04','特殊字符','🟢 低危','中文名/空格/符号','保存失败或报错'],
    ])

doc.add_page_break()

# ═════ 3. 怎么修 ═════
doc.add_heading('3. 怎么修——4 步安全加固', level=1)

doc.add_heading('3.1 修复 1：检查文件扩展名', level=2)
doc.add_paragraph('只允许特定的图片格式上传。')

p = doc.add_paragraph()
r = p.add_run('修复前：'); r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)
p.add_run(' 不检查，什么文件都能传')

p = doc.add_paragraph()
r = p.add_run('修复后：'); r.bold = True; r.font.color.rgb = RGBColor(0x00,0x80,0x00)
p.add_run(' 只允许 png/jpg/jpeg/gif/bmp/webp')

doc.add_paragraph(
    '# ✅ 代码\n'
    'ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "bmp", "webp"}\n\n'
    'def allowed_file(filename):\n'
    '    ext = filename.rsplit(".", 1)[-1].lower()\n'
    '    return ext in ALLOWED_EXTENSIONS'
)

doc.add_heading('3.2 修复 2：防路径穿越', level=2)
doc.add_paragraph('不管用户输入什么文件名，只取最后一部分（真正的文件名）。')

p = doc.add_paragraph()
r = p.add_run('修复前：'); r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)
p.add_run(' filename = f.filename  # 直接拿来用')

p = doc.add_paragraph()
r = p.add_run('修复后：'); r.bold = True; r.font.color.rgb = RGBColor(0x00,0x80,0x00)
p.add_run(' os.path.basename() 只取文件名')

doc.add_paragraph(
    '# ✅ 代码\n'
    'def safe_filename(filename):\n'
    '    # 去掉路径部分！防路径穿越\n'
    '    filename = os.path.basename(filename)\n'
    '    # 去掉危险字符\n'
    '    filename = re.sub(r\'[^\\w\\.\\-]\', \'_\', filename)\n'
    '    return filename'
)

p = doc.add_paragraph()
r = p.add_run('原理：'); r.bold = True
p.add_run('用户传 "../../shell.php" → basename 处理后只剩 "shell.php"，../ 全被去掉 ✅')

doc.add_heading('3.3 修复 3：同名文件自动重命名', level=2)
doc.add_paragraph('如果文件已存在，自动在文件名后加一串随机字符，避免覆盖。')

doc.add_paragraph(
    '# ✅ 代码\n'
    'if os.path.exists(save_path):\n'
    '    name, ext = os.path.splitext(filename)\n'
    '    filename = f"{name}_{secrets.token_hex(4)}{ext}"'
)

doc.add_heading('3.4 修复 4：大小校验前置', level=2)
doc.add_paragraph('先读文件内容检查大小，再保存，避免超大文件撑爆磁盘。')

doc.add_page_break()

# ═════ 4. 修复前后对比 ═════
doc.add_heading('4. 修复前后对比', level=1)
add_table(doc, ['对比项','修复前','修复后'],
    [
        ['文件类型','不检查，任意文件','仅限 png/jpg/gif/bmp/webp'],
        ['文件名','直接使用用户提交的名称','basename去路径 + 去特殊字符'],
        ['路径穿越','../../evil.exe 可写任意位置','os.path.basename 只保留文件名'],
        ['同名覆盖','后上传覆盖前一个','自动加随机后缀，不覆盖'],
        ['文件名长度','无限制','限制 100 字符以内'],
        ['文件大小','Flask 默认限制','先读取检查再保存'],
        ['日志','无记录','记录上传者和文件名'],
    ])

doc.add_page_break()

# ═════ 5. 新手总结 ═════
doc.add_heading('5. 新手总结', level=1)
tips = [
    ('永远不要相信用户上传的文件名', '坏人会把文件名改成 "../../病毒.exe" 来骗你'),
    ('文件类型一定要检查', '只允许你需要的格式，比如只让传图片'),
    ('路径穿越要防', '用 basename() 只取文件名，不要相信用户给的路径'),
    ('同名要处理', '加个随机数或者时间戳，避免互相覆盖'),
    ('大小要限制', '不然坏人给你传个 100GB 文件，服务器直接炸了'),
]
for title, desc in tips:
    p = doc.add_paragraph()
    r = p.add_run('✅ ' + title); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph(desc + '\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('记住：上传功能是网站最危险的功能之一，一定要好好保护！')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

output_path = '/home/user/projects/user-mgr/上传漏洞修复报告_新手版.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
print(f'📄 大小: {os.path.getsize(output_path) / 1024:.1f} KB')
