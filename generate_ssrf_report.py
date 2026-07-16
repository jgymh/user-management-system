#!/usr/bin/env python3
"""SSRF漏洞修复报告 - 新手入门版"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for level in range(1,4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'; hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, hc="2B579A"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), hc); c._tc.get_or_add_tcPr().append(shd)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
            if ri%2==0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "F0F4F8"); c._tc.get_or_add_tcPr().append(shd)
    return t

# 封面
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('SSRF漏洞修复报告'); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor(0x0A,0x1A,0x3A)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('NEXUS 用户管理系统 · 新手入门版'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━'*40).font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
for k,v in [('项目名称','NEXUS 用户管理系统'),('报告类型','SSRF漏洞修复（新手入门版）'),('仓库地址','https://github.com/jgymh/user-management-system')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{k}:  '); r.bold = True; r.font.size = Pt(11)
    p.add_run(v).font.size = Pt(11)
doc.add_page_break()

# 1. 什么是SSRF
doc.add_heading('1. 什么是 SSRF 漏洞', level=1)
doc.add_paragraph('SSRF（Server-Side Request Forgery）服务端请求伪造，简单说就是：')

p = doc.add_paragraph()
r = p.add_run('攻击者让服务器替他去访问他自己访问不到的地方。')
r.bold = True

doc.add_paragraph('用生活中的例子理解：')
doc.add_paragraph('你去餐厅点菜，你跟服务员说"我要一份红烧肉"（正常请求）。但如果你跟服务员说"你去厨房后门帮我看看冰箱里有啥"（SSRF攻击），服务员真的去帮你看了——这就是 SSRF。')

doc.add_paragraph()
doc.add_paragraph('攻击场景：')
p = doc.add_paragraph()
r = p.add_run('• 读取内网服务：')
r.bold = True
p.add_run('提交 http://127.0.0.1:5000，服务器访问自己的Flask服务')
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('• 读取云服务器元数据：')
r.bold = True
p.add_run('提交 http://169.254.169.254/latest/meta-data/（AWS/阿里云）')
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('• 扫描内网端口：')
r.bold = True
p.add_run('提交 http://192.168.1.1:22、http://10.0.0.1:3306 等')
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('• 读取本地文件：')
r.bold = True
p.add_run('提交 file:///etc/passwd，直接读取服务器文件')

doc.add_page_break()

# 2. 项目中发现的SSRF漏洞
doc.add_heading('2. 项目中发现了什么漏洞', level=1)

doc.add_heading('漏洞：/fetch-url 无任何限制', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🔴 严重')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

doc.add_paragraph('/fetch-url 接口直接使用 urllib.request.urlopen() 访问用户提交的URL，没有任何安全检查。')

p = doc.add_paragraph()
r = p.add_run('攻击方法 1：访问内网服务')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('POST /fetch-url  url=http://127.0.0.1:5000/')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 读取本机Flask应用的页面内容')

p = doc.add_paragraph()
r = p.add_run('攻击方法 2：扫描内网')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('POST /fetch-url  url=http://192.168.1.1:22')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 扫描内网SSH端口是否开放')

p = doc.add_paragraph()
r = p.add_run('攻击方法 3：读取本地文件')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('POST /fetch-url  url=file:///etc/passwd')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 读取服务器密码文件')

doc.add_paragraph()
doc.add_paragraph('修复前代码：')
p = doc.add_paragraph()
r = p.add_run('''# ❌ 漏洞代码
url = request.form.get("url", "")
req = urllib.request.Request(url)
resp = urllib.request.urlopen(req)  # 直接访问，不检查！''')
r.font.size = Pt(9); r.font.name = 'Courier New'

doc.add_paragraph('修复后代码：')
p = doc.add_paragraph()
r = p.add_run('''# ✅ 安全代码
parsed = urllib.parse.urlparse(url)
if parsed.scheme not in ("http", "https"):
    return "仅支持 http 和 https"

ip_addrs = socket.getaddrinfo(hostname, None)
for addr in ip_addrs:
    ip_obj = ipaddress.ip_address(addr[4][0])
    if ip_obj.is_private or ip_obj.is_loopback:
        return "禁止访问内网地址"''')
r.font.size = Pt(9); r.font.name = 'Courier New'; r.font.color.rgb = RGBColor(0x00,0x80,0x00)

doc.add_page_break()

# 3. 漏洞汇总表
doc.add_heading('3. 漏洞汇总表', level=1)
add_table(doc, ['编号','漏洞名称','风险','位置','攻击方式','后果'],
    [
        ['SEC-SSRF-01','内网服务访问','🔴 严重','/fetch-url','http://127.0.0.1','内网信息泄露'],
        ['SEC-SSRF-02','云元数据读取','🔴 严重','/fetch-url','http://169.254.169.254','云服务器密钥泄露'],
        ['SEC-SSRF-03','本地文件读取','🔴 严重','/fetch-url','file:///etc/passwd','系统文件泄露'],
        ['SEC-SSRF-04','内网端口扫描','🟠 高危','/fetch-url','http://10.0.0.1:3306','内网拓扑探测'],
    ])

doc.add_page_break()

# 4. 修复原理
doc.add_heading('4. 修复原理', level=1)

doc.add_heading('修复 1：协议白名单', level=2)
doc.add_paragraph('只允许 http:// 和 https:// 协议，拒绝 file://、dict://、gopher:// 等。')
add_table(doc, ['协议','修复前','修复后'],
    [['http://','✅ 允许','✅ 允许'],
     ['https://','✅ 允许','✅ 允许'],
     ['file://','✅ 允许','❌ 拒绝'],
     ['dict://','✅ 允许','❌ 拒绝'],
     ['gopher://','✅ 允许','❌ 拒绝']])

doc.add_heading('修复 2：DNS解析 + IP检测', level=2)
doc.add_paragraph('把域名解析成IP，然后检查IP是否为内网地址。Python的ipaddress模块内置了判断方法：')
doc.add_paragraph('• ip_obj.is_private → 192.168.x.x, 10.x.x.x, 172.16-31.x.x')
doc.add_paragraph('• ip_obj.is_loopback → 127.0.0.1, ::1')
doc.add_paragraph('• ip_obj.is_link_local → 169.254.x.x')

doc.add_heading('修复 3：绕过防护测试', level=2)
doc.add_paragraph('即使是使用了内网域名（如 localhost），DNS解析后得到127.0.0.1，也会被拦截。')

doc.add_page_break()

# 5. 修复前后对比
doc.add_heading('5. 修复前后对比', level=1)
add_table(doc, ['检测项','修复前','修复后'],
    [
        ['访问 127.0.0.1','✅ 可访问','❌ 禁止内网'],
        ['访问 192.168.x.x','✅ 可访问','❌ 禁止内网'],
        ['访问 10.x.x.x','✅ 可访问','❌ 禁止内网'],
        ['读取 file://','✅ 可读取','❌ 协议拒绝'],
        ['访问 169.254.x.x','✅ 可访问','❌ 禁止链路本地'],
        ['访问外网(baidu)','✅ 正常','✅ 正常'],
    ], hc="1A6B3C")

doc.add_page_break()

# 6. 新手总结
doc.add_heading('6. 新手总结', level=1)

tips = [
    ('永远不要直接使用用户输入的URL', '一定要先解析URL，检查协议和主机地址'),
    ('协议白名单是最基本的', '只允许 http/https，拒绝其他一切协议'),
    ('内网IP必须阻止', '127.0.0.1、10.x.x.x、172.16-31.x.x、192.168.x.x、169.254.x.x 全都要拦截'),
    ('DNS绕过要防', '用户可能用 localhost、internal.api 等域名指向内网，需要DNS解析后检查IP'),
    ('可以利用 Python 内置库', 'urllib.parse 解析URL、socket 做DNS解析、ipaddress 判断IP类型'),
]
for title, desc in tips:
    p = doc.add_paragraph()
    r = p.add_run('✅ ' + title); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph(desc + '\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('记住：服务器发出去的请求也要管！用户让服务器去访问的东西，不一定都是安全的！')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

output_path = '/home/user/projects/user-mgr/SSRF漏洞修复报告_新手版.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
