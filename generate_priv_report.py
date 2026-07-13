#!/usr/bin/env python3
"""权限漏洞修复报告 - 新手入门版"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for level in range(1,4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'; hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, hc="2B579A"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), hc); c._tc.get_or_add_tcPr().append(shd)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
            if ri%2==0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "F0F4F8"); c._tc.get_or_add_tcPr().append(shd)
    return t

# ═══════ 封面 ═══════
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('权限漏洞修复报告'); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor(0x0A,0x1A,0x3A)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('NEXUS 用户管理系统 · 新手入门版'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━'*40).font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
for k,v in [('项目名称','NEXUS 用户管理系统'),('报告类型','权限漏洞修复（新手入门版）'),('仓库地址','https://github.com/jgymh/user-management-system')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{k}:  '); r.bold = True; r.font.size = Pt(11)
    p.add_run(v).font.size = Pt(11)
doc.add_page_break()

# ═════ 1. 什么是权限漏洞 ═════
doc.add_heading('1. 什么是权限漏洞', level=1)
doc.add_paragraph('权限漏洞就是一个用户能干"他不该干"的事情。')
p = doc.add_paragraph()
r = p.add_run('正常情况：')
r.bold = True
p.add_run('你登录自己的账户，只能看自己的资料，花自己的钱。')

p = doc.add_paragraph()
r = p.add_run('有漏洞时：')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)
p.add_run('你不仅能看自己的，改个ID就能看别人的；还能给别人的账号充钱或扣钱。')

doc.add_paragraph('用简单的例子理解：就像你去银行取钱，柜员只看你递过去的卡号，不看你身份证——你拿别人的卡也能取钱。')

doc.add_page_break()

# ═════ 2. 发现什么漏洞 ═════
doc.add_heading('2. 项目中发现了哪些权限漏洞', level=1)

doc.add_heading('漏洞 1：越权查看他人资料（IDOR）', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🔴 严重')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

doc.add_paragraph('/profile?user_id=1 看 admin，改成 ?user_id=2 就能看 alice 的资料。系统没有验证"你是否有权看这个人"。')

p = doc.add_paragraph()
r = p.add_run('攻击方法：')
r.bold = True
p.add_run('在浏览器地址栏直接把 user_id=1 改成 user_id=2')

doc.add_paragraph('修复前代码：')
p = doc.add_paragraph()
r = p.add_run('''# ❌ 漏洞代码
user_id = request.args.get("user_id", "")  # 从URL获取ID
c.execute("SELECT ... FROM users WHERE id=?", (user_id,))
# 没有检查当前登录用户是否等于这个ID！''')
r.font.size = Pt(9); r.font.name = 'Courier New'

doc.add_paragraph('修复后代码：')
p = doc.add_paragraph()
r = p.add_run('''# ✅ 安全代码
username = session.get("username")  # 从session获取当前用户
c.execute("SELECT ... FROM users WHERE username=?", (username,))
# 不管URL传什么user_id，只看session里的''')
r.font.size = Pt(9); r.font.name = 'Courier New'; r.font.color.rgb = RGBColor(0x00,0x80,0x00)

doc.add_heading('漏洞 2：越权充值（修改他人余额）', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🔴 严重')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

doc.add_paragraph('充值接口接收表单里的 user_id，可以给任何人充值或扣钱。')

p = doc.add_paragraph()
r = p.add_run('攻击方法：')
r.bold = True
p.add_run('充值时抓包修改 user_id 参数，就能给别人的账号充钱或扣钱。')

doc.add_paragraph('修复前代码：')
p = doc.add_paragraph()
r = p.add_run('''# ❌ 漏洞代码
user_id = request.form.get("user_id", "")  # 从表单获取
amount_num = int(amount)  # 不做正负判断
c.execute("UPDATE users SET balance = balance + ? WHERE id=?", (amount_num, user_id))
# 给谁都行！负数也行！''')
r.font.size = Pt(9); r.font.name = 'Courier New'

doc.add_paragraph('修复后代码：')
p = doc.add_paragraph()
r = p.add_run('''# ✅ 安全代码
username = session.get("username")  # 从session获取
if amount_num <= 0:  # 金额必须大于0
    return "充值金额必须大于0"
c.execute("UPDATE users SET balance = balance + ? WHERE username=?", (amount_num, username))
# 只能给自己充！不能充负数！''')
r.font.size = Pt(9); r.font.name = 'Courier New'; r.font.color.rgb = RGBColor(0x00,0x80,0x00)

doc.add_heading('漏洞 3：未登录访问', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🟠 高危')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x66,0x00)

doc.add_paragraph('不登录也能直接访问 /profile 和 /recharge，相当于任何人都能操作。')
doc.add_paragraph('修复方式：在每个路由开头加一句 if "username" not in session: return redirect("/login")')

doc.add_heading('漏洞 4：无 CSRF 保护的充值', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🟠 高危')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x66,0x00)

doc.add_paragraph('充值接口没有 CSRF Token 校验。攻击者可构造恶意页面，让用户浏览器自动提交充值请求，给攻击者账号充钱。')
doc.add_paragraph('修复方式：充值接口增加了 CSRF Token 校验，和登录接口一样。')

doc.add_page_break()

# ═════ 3. 漏洞汇总表 ═════
doc.add_heading('3. 漏洞汇总表', level=1)
add_table(doc, ['编号','漏洞名称','风险','位置','攻击方式','后果'],
    [
        ['SEC-PRIV-01','越权查看资料','🔴 严重','/profile','修改URL的user_id','任意用户资料泄露'],
        ['SEC-PRIV-02','越权充值','🔴 严重','/recharge','修改表单user_id','给任意账号加/扣钱'],
        ['SEC-PRIV-03','未登录访问','🟠 高危','/profile /recharge','直接访问URL','未授权操作'],
        ['SEC-PRIV-04','CSRF充值','🟠 高危','/recharge','构造恶意表单','替他人充值'],
        ['SEC-PRIV-05','负数充值','🔴 严重','/recharge','amount=-99999','无限扣钱'],
    ])

doc.add_page_break()

# ═════ 4. 修复前后对比 ═════
doc.add_heading('4. 修复前后对比', level=1)
add_table(doc, ['检查项','修复前','修复后'],
    [
        ['未登录访问/profile','✅ 可访问','❌ 跳转登录页'],
        ['查看别人资料','✅ user_id=2 看 alice','❌ 只能看自己的'],
        ['未登录充值','✅ 可充值','❌ 跳转登录页'],
        ['给他人充值','✅ 改user_id就行','❌ 只能给自己充'],
        ['充负数扣钱','✅ amount=-999','❌ 拦截并报错'],
        ['CSRF充值防御','❌ 无','✅ 有Token校验'],
    ], hc="1A6B3C")

doc.add_page_break()

# ═════ 5. 新手总结 ═════
doc.add_heading('5. 新手总结', level=1)

tips = [
    ('永远不要相信前端传来的 user_id', '要从 session 里取当前登录用户的信息'),
    ('该登录的地方必须检查登录', '每个涉及用户数据的接口都要校验 session'),
    ('金额操作必须校验正负', '不校验的话-99999就是无限提款机'),
    ('重要操作要加 CSRF 防护', '充值和登录一样重要，必须有 Token 校验'),
    ('记住一句话', '"用户传什么就信什么" = 漏洞。 "从 session 拿" = 安全'),
]
for title, desc in tips:
    p = doc.add_paragraph()
    r = p.add_run('✅ ' + title); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph(desc + '\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('核心原则：永远从 session 获取当前用户身份，不要相信前端传来的用户ID！')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

output_path = '/home/user/projects/user-mgr/权限漏洞修复报告_新手版.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
print(f'📄 大小: {os.path.getsize(output_path) / 1024:.1f} KB')
