#!/usr/bin/env python3
"""文件包含漏洞修复报告 - 新手入门版"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for level in range(1,4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'; hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, hc="2B579A"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), hc); c._tc.get_or_add_tcPr().append(shd)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
            if ri%2==0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "F0F4F8"); c._tc.get_or_add_tcPr().append(shd)
    return t

# ═══════ 封面 ═══════
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('文件包含漏洞修复报告'); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor(0x0A,0x1A,0x3A)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('NEXUS 用户管理系统 · 新手入门版'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━'*40).font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
for k,v in [('项目名称','NEXUS 用户管理系统'),('报告类型','文件包含漏洞修复（新手入门版）'),('仓库地址','https://github.com/jgymh/user-management-system')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{k}:  '); r.bold = True; r.font.size = Pt(11)
    p.add_run(v).font.size = Pt(11)
doc.add_page_break()

# ═════ 1. 什么是文件包含漏洞 ═════
doc.add_heading('1. 什么是文件包含漏洞', level=1)
doc.add_paragraph('文件包含漏洞（File Inclusion）就是网站允许用户指定"加载哪个文件"，但又没做安全检查，导致攻击者可以读取服务器上的任意文件。')

p = doc.add_paragraph()
r = p.add_run('用生活中的例子理解：')
r.bold = True
doc.add_paragraph('你去图书馆借书，正常流程是告诉管理员"我要借 help.html"。管理员去「pages」书架上拿给你。')
doc.add_paragraph('但如果有漏洞，你可以说"我要借 ../app.py"（跑到书架外面去了），管理员也真的跑出去帮你把 app.py 拿来了。')
p = doc.add_paragraph()
r = p.add_run('更严重的："我要借 ../../../../etc/passwd"')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)
doc.add_paragraph('管理员一层层往上翻，把 Linux 系统的密码文件都拿出来了。这就是路径穿越（Path Traversal）。')

doc.add_page_break()

# ═════ 2. 项目中发现的漏洞 ═════
doc.add_heading('2. 项目中发现了什么漏洞', level=1)

doc.add_heading('漏洞：/page 路由路径穿越', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🔴 严重')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

doc.add_paragraph('我们的系统有个 /page 路由，可以根据参数 name 加载 pages/ 目录下的文件。')
doc.add_paragraph('但代码直接把用户输入的 name 拼接到路径里，没有做任何过滤。')

p = doc.add_paragraph()
r = p.add_run('攻击方法 1：读取源码')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('/page?name=../app.py')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 读取到 Flask 应用的全部源码，包括 secret_key、数据库配置等。')

p = doc.add_paragraph()
r = p.add_run('攻击方法 2：读取系统密码文件')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('/page?name=../../../../etc/passwd')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 读取 Linux 系统的 /etc/passwd 文件，获取所有用户名。')

doc.add_paragraph('修复前代码：')
p = doc.add_paragraph()
r = p.add_run('''# ❌ 漏洞代码
name = request.args.get("name", "")
page_path = os.path.join("pages", name)
# 直接拼接，233.. 也会被拼进去！
# pages/../app.py = 读取 app.py''')
r.font.size = Pt(9); r.font.name = 'Courier New'

doc.add_paragraph('修复后代码：')
p = doc.add_paragraph()
r = p.add_run('''# ✅ 安全代码
name = request.args.get("name", "")
name = os.path.basename(name)  # 只取文件名，去掉路径！
page_path = os.path.join("pages", name)
# 传 ../app.py → basename 后只剩 app.py
# pages/app.py 不存在 → "页面不存在"''')
r.font.size = Pt(9); r.font.name = 'Courier New'; r.font.color.rgb = RGBColor(0x00,0x80,0x00)

doc.add_page_break()

# ═════ 3. 漏洞汇总表 ═════
doc.add_heading('3. 漏洞汇总表', level=1)
add_table(doc, ['编号','漏洞名称','风险','位置','攻击方式','后果'],
    [
        ['SEC-FILE-01','路径穿越读取源码','🔴 严重','/page?name=','../app.py','源码泄露(密钥/配置)'],
        ['SEC-FILE-02','路径穿越读取系统文件','🔴 严重','/page?name=','../../etc/passwd','系统用户信息泄露'],
        ['SEC-FILE-03','任意文件读取','🔴 严重','/page?name=','任意../组合','读取服务器任意文件'],
    ])

doc.add_page_break()

# ═════ 4. 修复原理 ═════
doc.add_heading('4. 修复原理：os.path.basename', level=1)
doc.add_paragraph('修复的核心就是一行代码：name = os.path.basename(name)')

doc.add_paragraph('basename 的作用是"只取路径中的文件名部分，去掉所有目录部分"。')

add_table(doc, ['用户输入','os.path.basename 处理结果','拼接后路径','结果'],
    [
        ['help','help','pages/help ✅','正常读取'],
        ['../app.py','app.py','pages/app.py ❌','找不到文件'],
        ['../../etc/passwd','passwd','pages/passwd ❌','找不到文件'],
        ['../app.py%00','app.py%00','pages/app.py%00 ❌','找不到文件'],
    ])

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('核心原理一句话：')
r.bold = True
p.add_run('不管用户怎么加 ../，basename 只保留最右边的文件名，路径部分全部丢弃。')

doc.add_page_break()

# ═════ 5. 修复前后对比 ═════
doc.add_heading('5. 修复前后对比', level=1)
add_table(doc, ['检测项','修复前','修复后'],
    [
        ['正常访问 help','✅ 显示帮助页面','✅ 显示帮助页面'],
        ['../app.py 读取源码','✅ 读到 Flask 源码','❌ 页面不存在'],
        ['../../etc/passwd','✅ 读到系统密码文件','❌ 页面不存在'],
        ['任意路径穿越','✅ 可读取任意文件','❌ 全部拦截'],
    ], hc="1A6B3C")

doc.add_page_break()

# ═════ 6. 新手总结 ═════
doc.add_heading('6. 新手总结', level=1)

tips = [
    ('basename 是你的好朋友', '只要遇到"用户传文件名"的场景，先用 basename 过滤一下'),
    ('永远不要拼接用户输入到路径', 'name = request.args.get("name") + os.path.join("pages", name) 中间必须加过滤'),
    ('不是只有 ../ 才危险', 'url编码 %2e%2e%2f、unicode 编码等都可以绕过关键字过滤，但 basename 全部防住'),
    ('文件包含比 SQL 注入还隐蔽', 'SQL 注入会报错让你发现，文件包含悄无声息地就把源码偷走了'),
]
for title, desc in tips:
    p = doc.add_paragraph()
    r = p.add_run('✅ ' + title); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph(desc + '\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('记住：文件路径操作前，先过 basename！')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

output_path = '/home/user/projects/user-mgr/文件包含漏洞修复报告_新手版.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
