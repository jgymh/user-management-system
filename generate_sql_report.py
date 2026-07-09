#!/usr/bin/env python3
"""生成 SQL 注入防护报告 — 新手小白版"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 全局字体 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def set_cell_shading(cell, color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shd)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(cell, "2B579A")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if ri % 2 == 0:
                set_cell_shading(cell, "F0F4F8")
    return table

# ════════════════════════════════════════
# 封面
# ════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SQL 注入防护报告')
run.font.size = Pt(30)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0A, 0x1A, 0x3A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('NEXUS 用户管理系统 · 新手入门版')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph()

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_line.add_run('━' * 40).font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph()

for k, v in [
    ('项目名称', 'NEXUS 用户管理系统'),
    ('报告类型', 'SQL 注入防护（新手入门版）'),
    ('仓库地址', 'https://github.com/jgymh/user-management-system'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{k}:  ')
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(v).font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════
# 1. 什么是 SQL 注入
# ════════════════════════════════════════
doc.add_heading('1. 什么是 SQL 注入', level=1)

doc.add_paragraph(
    'SQL 注入是 Web 安全中最常见、最危险的漏洞之一。简单来说就是：'
    '攻击者在输入框或网址参数里"夹带"SQL 代码，让数据库执行不该执行的命令。'
)

doc.add_paragraph('用一个生活中的例子来理解：')
p = doc.add_paragraph()
p.add_run('假设你是学校保安，进出校门要报学号。')
r = p.add_run('正常学生说"2024001"——你查表，放行。')
r.bold = True
doc.add_paragraph(
    '但有一天来了个人说："2024001 OR 1=1"，意思是"学号是2024001或者1=1"。'
    '因为1=1永远成立，结果你把所有人都放进去了！这就是 SQL 注入。'
)

doc.add_paragraph(
    '在我们的项目里，搜索功能就遇到过这个问题。'
    '用户在搜索框输入 SQL 代码，程序直接把代码拼到 SQL 查询里，数据库就执行了。'
)

doc.add_page_break()

# ════════════════════════════════════════
# 2. 项目中发现的 3 个漏洞
# ════════════════════════════════════════
doc.add_heading('2. 项目里发现了哪几个漏洞', level=1)
doc.add_paragraph('在我们的 NEXUS 用户管理系统中，一共发现了 3 处 SQL 注入漏洞，都在 app.py 文件里。')

doc.add_heading('漏洞 1：搜索功能（首页搜索框）', level=2)
doc.add_paragraph(
    '用户在搜索框输入关键词后，程序用 f-string 把关键词直接拼进了 SQL 语句。'
    '攻击者在搜索框输入特殊字符，就能让 SQL 执行额外的查询。'
)

p = doc.add_paragraph()
r = p.add_run('漏洞代码位置：')
r.bold = True
p.add_run('app.py 第 190 行（修复前）')

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('原始代码：')
r.bold = True
p.add_run('''sql = f"SELECT * FROM users WHERE username LIKE '%{keyword}%'"''')

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('攻击方式：')
r.bold = True
p.add_run('搜索框输入  →  \' UNION SELECT 1,username,password,phone FROM users --')

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('攻击效果：')
r.bold = True
p.add_run('所有用户的密码被窃取，显示在搜索结果中')

doc.add_heading('漏洞 2：搜索接口（/search）', level=2)
doc.add_paragraph(
    '跟首页搜索一样的问题，/search 接口也是直接用 f-string 拼接 SQL。'
    '攻击者通过网址参数就能发起注入攻击。'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('攻击方式：')
r.bold = True
p.add_run('curl "http://.../search?keyword=\' OR \'1\'=\'1"')

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('攻击效果：')
r.bold = True
p.add_run('一次搜索返回数据库所有用户的信息')

doc.add_heading('漏洞 3：注册功能（/register）', level=2)
doc.add_paragraph(
    '注册时用户名填入 SQL 代码，同样因为 f-string 拼接，可以篡改数据库。'
    '甚至可以把管理员密码改掉。'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('攻击方式：')
r.bold = True
p.add_run('注册时用户名填 → admin\', \'hacked\', \'x@x.com\', \'000\') ON CONFLICT(username) DO UPDATE SET password=\'hacked\' --')

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('攻击效果：')
r.bold = True
p.add_run('管理员密码被篡改，原主人无法登录')

doc.add_page_break()

# ════════════════════════════════════════
# 3. 漏洞汇总表
# ════════════════════════════════════════
doc.add_heading('3. 漏洞汇总表', level=1)

add_table(doc,
    ['编号', '漏洞位置', '风险等级', '攻击方式', '攻击后果'],
    [
        ['SEC-SQL-01', '/ 首页搜索', '🔴 严重', '搜索框输入 SQL 代码', '窃取所有用户密码'],
        ['SEC-SQL-02', '/search 接口', '🔴 严重', 'URL 参数注入', '获取全部用户信息'],
        ['SEC-SQL-03', '/register 注册', '🔴 严重', '用户名注入 SQL', '篡改管理员密码'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════
# 4. 最基础的修复方法：参数化查询
# ════════════════════════════════════════
doc.add_heading('4. 怎么修 —— 参数化查询', level=1)

doc.add_paragraph(
    '修复 SQL 注入最基础、最有效的方法就是"参数化查询"。'
    '原理很简单：把 SQL 语句和数据分开，让数据库知道"哪部分是命令，哪部分是数据"。'
)

doc.add_paragraph(
    '在我们项目里用的是 Python 的 sqlite3 模块，修复方式就是用 ? 占位符替换 f-string。'
)

doc.add_heading('4.1 修复前（有漏洞）', level=2)

p = doc.add_paragraph()
r = p.add_run('错误写法（f-string 拼接）：')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph(
    '# ❌ 漏洞代码：用户输入直接拼进 SQL\n'
    'keyword = request.args.get("keyword", "")\n'
    'sql = f"SELECT * FROM users WHERE username LIKE \'%{keyword}%\'"\n'
    'c.execute(sql)'
)

doc.add_paragraph(
    '这段代码的问题：keyword 的值如果是 \' OR \'1\'=\'1，SQL 就变成了：\n'
    "SELECT * FROM users WHERE username LIKE '%' OR '1'='1%'\n"
    'OR \'1\'=\'1 永远为真 → 返回全部数据'
)

doc.add_heading('4.2 修复后（安全）', level=2)

p = doc.add_paragraph()
r = p.add_run('正确写法（? 占位符）：')
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

doc.add_paragraph(
    '# ✅ 安全代码：数据和 SQL 分开\n'
    'keyword = request.args.get("keyword", "")\n'
    'sql = "SELECT * FROM users WHERE username LIKE ?"\n'
    'c.execute(sql, (\'%\' + keyword + \'%\',))'
)

doc.add_paragraph(
    '改成 ? 后，keyword 的值只会被当作"数据"传入，不会被当作 SQL 命令执行。'
    '即使搜索框输入 \' OR \'1\'=\'1，数据库也只会把它当成普通的搜索文字，'
    '去查找用户名里包含 "\\\' OR \\\'1\\\'=\\\'1" 的用户。'
)

doc.add_heading('4.3 修复前后对比', level=2)

add_table(doc,
    ['对比项', '修复前（f-string 拼接）', '修复后（? 参数化）'],
    [
        ['SQL 拼接方式', 'f"WHERE name LIKE " + keyword', "WHERE name LIKE ?"],
        ['用户输入', "直接插入 SQL 语句", "作为参数传入，不改变 SQL"],
        ["输入 ' OR 1=1", "WHERE条件被篡改 → 全表返回", "搜索文字本身 → 找不到任何用户"],
        ["输入 UNION SELECT", "UNION查询被执行，窃取密码", "搜索文字本身，不会执行额外查询"],
        ['安全性', '❌ 高危', '✅ 安全'],
    ])

doc.add_page_break()

# ════════════════════════════════════════
# 5. 项目代码中的具体修复
# ════════════════════════════════════════
doc.add_heading('5. 项目代码里的具体修复', level=1)

doc.add_paragraph(
    '以下是我们项目 app.py 中 3 处漏洞的具体修复过程。'
)

doc.add_heading('5.1 修复首页搜索（第 190 行）', level=2)

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('修复前：')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run(' sql = f"SELECT ... WHERE username LIKE \'%{keyword}%\' OR email LIKE \'%{keyword}%\'"')

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('修复后：')
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run(' sql = "SELECT ... WHERE username LIKE ? OR email LIKE ?"')
doc.add_paragraph(' c.execute(sql, (\'%\' + keyword + \'%\', \'%\' + keyword + \'%\'))', style='List Bullet')

doc.add_heading('5.2 修复搜索接口 /search（第 317 行）', level=2)

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('修复前：')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run(' 同样用 f-string 拼接 keyword')

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('修复后：')
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run(' 改成 ? 占位符传参，跟首页搜索一样')

doc.add_heading('5.3 修复注册接口 /register（第 294 行）', level=2)

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('修复前：')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run(' sql = f"INSERT INTO users VALUES (\'{username}\', \'{password}\', ...)"')

p = doc.add_paragraph()
p.style = 'List Bullet'
r = p.add_run('修复后：')
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run(' sql = "INSERT INTO users VALUES (?, ?, ?, ?)"')
doc.add_paragraph(' c.execute(sql, (username, password, email, phone))', style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════
# 6. 验证修复效果
# ════════════════════════════════════════
doc.add_heading('6. 验证修复有没有用', level=1)

doc.add_paragraph(
    '修复完之后，用同样的攻击方法再试一次，看还能不能成功。'
)

doc.add_heading('6.1 测试 UNION 注入', level=2)

p = doc.add_paragraph()
r = p.add_run('修复前：')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run(' 搜索框输入 UNION 查询 → 密码全部显示在页面上 ✅')

p = doc.add_paragraph()
r = p.add_run('修复后：')
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run(' 搜索框输入同样的 UNION 查询 → 不显示任何密码 ❌')

doc.add_heading('6.2 测试 OR 注入', level=2)

p = doc.add_paragraph()
r = p.add_run('修复前：')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run(' 输入 \' OR \'1\'=\'1 → 返回所有用户 ✅')

p = doc.add_paragraph()
r = p.add_run('修复后：')
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run(' 输入同样的内容 → 搜索不到任何用户 ❌')

doc.add_heading('6.3 测试注册注入', level=2)

p = doc.add_paragraph()
r = p.add_run('修复前：')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run(' 注册时用户名注入 ON CONFLICT → 数据库里 admin 密码被篡改 ✅')

p = doc.add_paragraph()
r = p.add_run('修复后：')
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.add_run(' 同样的用户名 → 正常注册搜索文字，不会篡改数据 ❌')

doc.add_page_break()

# ════════════════════════════════════════
# 7. 新手总结
# ════════════════════════════════════════
doc.add_heading('7. 新手总结', level=1)

doc.add_paragraph(
    '对于刚开始学安全的新手，记住这几点就够了：'
)

tips = [
    ('永远不要用 f-string 拼接 SQL', '这是 SQL 注入的根源。用户输入的数据可能包含恶意 SQL 代码。'),
    ('用 ? 占位符代替 %s', '参数化查询让数据库区分"命令"和"数据"，是最简单有效的防护。'),
    ('不要信任用户输入', '不管是搜索框、网址参数、还是注册表单，用户的任何输入都可能是攻击。'),
    ('修完后要验证', '用同样的攻击方法再试一次，确保真的修好了。'),
]

for title, desc in tips:
    p = doc.add_paragraph()
    r = p.add_run('✅ ' + title)
    r.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph(desc)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('记住一条规则：用户输入永远不要直接拼进 SQL！')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━' * 30).font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('报告结束').font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── 保存 ──
output_path = '/home/user/projects/user-mgr/SQL注入防护报告_新手版.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
print(f'📄 大小: {os.path.getsize(output_path) / 1024:.1f} KB')
