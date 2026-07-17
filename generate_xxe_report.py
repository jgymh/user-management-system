#!/usr/bin/env python3
"""XXE漏洞修复报告 - 新手入门版"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for level in range(1,4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'; hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, hc="2B579A"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), hc); c._tc.get_or_add_tcPr().append(shd)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
            if ri%2==0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "F0F4F8"); c._tc.get_or_add_tcPr().append(shd)
    return t

# 封面
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('XXE漏洞修复报告'); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor(0x0A,0x1A,0x3A)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('NEXUS 用户管理系统 · 新手入门版'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━'*40).font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
for k,v in [('项目名称','NEXUS 用户管理系统'),('报告类型','XXE漏洞修复（新手入门版）'),('仓库地址','https://github.com/jgymh/user-management-system')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{k}:  '); r.bold = True; r.font.size = Pt(11)
    p.add_run(v).font.size = Pt(11)
doc.add_page_break()

# 1. 什么是XXE
doc.add_heading('1. 什么是 XXE 漏洞', level=1)
doc.add_paragraph('XXE（XML External Entity）XML外部实体注入，简单说就是：攻击者在XML里夹带"外部实体"定义，让服务器读取本地文件或访问内网。')

p = doc.add_paragraph()
r = p.add_run('用生活中的例子理解：')
r.bold = True
doc.add_paragraph('你填了一张表格（XML），在表格里写了一行小字"请去隔壁档案室把机密文件拿来"（ENTITY定义）。办事员（XML解析器）真的去拿了——这就是XXE。')

doc.add_paragraph()
doc.add_paragraph('攻击原理：')
p = doc.add_paragraph()
r = p.add_run('正常XML：')
r.bold = True
p.add_run(' <name>张三</name>')
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('恶意XML：')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)
doc.add_paragraph('''<?xml version="1.0"?>
<!DOCTYPE foo [<!ENTITY xxe SYSTEM "/etc/passwd">]>
<name>&xxe;</name>''')
doc.add_paragraph('→ &xxe; 被替换为 /etc/passwd 的文件内容')

doc.add_page_break()

# 2. 项目中发现的漏洞
doc.add_heading('2. 项目中发现了什么漏洞', level=1)

doc.add_heading('漏洞：/xml-import 接口XXE注入', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🔴 严重')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

doc.add_paragraph('/xml-import 功能会自动检测用户XML中的 <!ENTITY 定义，提取文件路径并读取本地文件内容。')

p = doc.add_paragraph()
r = p.add_run('攻击方法 1：读取系统文件')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('''<!ENTITY xxe SYSTEM "/etc/passwd"> → <name>&xxe;</name>''')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 读取Linux系统的所有用户账号信息')

p = doc.add_paragraph()
r = p.add_run('攻击方法 2：读取源码')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('''<!ENTITY xxe SYSTEM "/home/.../app.py"> → <name>&xxe;</name>''')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 读取Flask应用源代码，包括secret_key、数据库配置')

p = doc.add_paragraph()
r = p.add_run('攻击方法 3：读取配置文件')
r.bold = True
p = doc.add_paragraph()
r = p.add_run('''<!ENTITY xxe SYSTEM "/etc/nginx/nginx.conf">''')
r.font.name = 'Courier New'; r.font.size = Pt(10)
doc.add_paragraph('→ 读取Nginx配置，探测内网反向代理设置')

doc.add_paragraph()
doc.add_paragraph('修复前代码：')
p = doc.add_paragraph()
r = p.add_run('''# ❌ 漏洞代码
entity_matches = re.findall(r'<!ENTITY.*SYSTEM."([^"]+)"', xml)
for filepath in entity_matches:
    with open(filepath, "r") as f:        # 直接读取！
        content = f.read()
        xml = xml.replace("&xxe;", content) # 替换实体！''')
r.font.size = Pt(9); r.font.name = 'Courier New'

doc.add_paragraph('修复后代码：')
p = doc.add_paragraph()
r = p.add_run('''# ✅ 安全代码
# 移除DTD声明，禁用外部实体
safe_xml = re.sub(r'<!DOCTYPE[^>]*>', '', xml_data)
safe_xml = re.sub(r'&[^;\\s]+;', '', safe_xml)  # 移除实体引用
root = ET.fromstring(safe_xml)  # 安全解析''')
r.font.size = Pt(9); r.font.name = 'Courier New'; r.font.color.rgb = RGBColor(0x00,0x80,0x00)

doc.add_page_break()

# 3. 漏洞汇总表
doc.add_heading('3. 漏洞汇总表', level=1)
add_table(doc, ['编号','漏洞名称','风险','位置','攻击方式','后果'],
    [
        ['SEC-XXE-01','读取系统文件','🔴 严重','/xml-import','ENTITY /etc/passwd','系统用户信息泄露'],
        ['SEC-XXE-02','读取源码','🔴 严重','/xml-import','ENTITY app.py','源代码泄露'],
        ['SEC-XXE-03','读取配置文件','🔴 严重','/xml-import','ENTITY nginx.conf','配置信息泄露'],
        ['SEC-XXE-04','SSRF+XXE组合','🔴 严重','/xml-import','ENTITY http://内网IP','内网探测'],
    ])

doc.add_page_break()

# 4. 修复原理
doc.add_heading('4. 修复原理', level=1)

doc.add_heading('修复 1：移除DTD声明', level=2)
doc.add_paragraph('DTD（Document Type Definition）是XML中定义实体的地方。去掉 <!DOCTYPE ...> 标签，实体定义就没了。')
p = doc.add_paragraph()
r = p.add_run('re.sub(r\'<!DOCTYPE[^>]*>\', \'\', xml_data)')
r.font.size = Pt(10); r.font.name = 'Courier New'

doc.add_heading('修复 2：移除实体引用', level=2)
doc.add_paragraph('即使DTD被去掉了，如果XML里还有 &xxe; 这样的实体引用，解析时会报错。所以把 &xxx; 格式的实体引用也清掉。')
p = doc.add_paragraph()
r = p.add_run('re.sub(r\'&[^;\\s]+;\', \'\', xml_data)')
r.font.size = Pt(10); r.font.name = 'Courier New'

doc.add_heading('修复 3：使用标准XML解析器', level=2)
doc.add_paragraph('Python的 xml.etree.ElementTree 默认不加载外部实体，但前提是DTD已经被移除。两步配合就安全了。')

doc.add_page_break()

# 5. 修复前后对比
doc.add_heading('5. 修复前后对比', level=1)
add_table(doc, ['检测项','修复前','修复后'],
    [
        ['正常XML解析','✅ 正常','✅ 正常'],
        ['XXE读/etc/passwd','✅ 读到root信息 ❌','❌ 解析失败 ✅'],
        ['XXE读app.py源码','✅ 读到Flask代码 ❌','❌ 解析失败 ✅'],
        ['处理DTD声明','解析并执行','移除'],
        ['处理实体引用','读取文件替换','清空'],
    ], hc="1A6B3C")

doc.add_page_break()

# 6. 新手总结
doc.add_heading('6. 新手总结', level=1)
tips = [
    ('XML解析时必须禁用外部实体', 'DTD里的 SYSTEM 关键字就是用来引用外部文件的'),
    ('re.sub()是你的好帮手', '用正则表达式把 DTD 和实体引用提前干掉，比配置解析器参数简单'),
    ('永远不要相信用户提交的XML', '用户提交的XML可能包含恶意实体定义，绝对不能直接解析'),
    ('测试XXE三板斧', '/etc/passwd 读不到 = 修复成功；app.py 读不到 = 修复成功'),
]
for title, desc in tips:
    p = doc.add_paragraph()
    r = p.add_run('✅ ' + title); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph(desc + '\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('记住：处理XML之前，先把DTD和实体干掉！')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

output_path = '/home/user/projects/user-mgr/XXE漏洞修复报告_新手版.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
