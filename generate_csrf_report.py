#!/usr/bin/env python3
"""CSRF漏洞修复报告 - 新手入门版"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for level in range(1,4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'; hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, hc="2B579A"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), hc); c._tc.get_or_add_tcPr().append(shd)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
            if ri%2==0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "F0F4F8"); c._tc.get_or_add_tcPr().append(shd)
    return t

# 封面
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CSRF漏洞修复报告'); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor(0x0A,0x1A,0x3A)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('NEXUS 用户管理系统 · 新手入门版'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('━'*40).font.color.rgb = RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph()
for k,v in [('项目名称','NEXUS 用户管理系统'),('报告类型','CSRF漏洞修复（新手入门版）'),('仓库地址','https://github.com/jgymh/user-management-system')]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{k}:  '); r.bold = True; r.font.size = Pt(11)
    p.add_run(v).font.size = Pt(11)
doc.add_page_break()

# 1. 什么是CSRF
doc.add_heading('1. 什么是 CSRF 漏洞', level=1)
doc.add_paragraph('CSRF（Cross-Site Request Forgery）跨站请求伪造，简单说就是：')

p = doc.add_paragraph()
r = p.add_run('你在 A 网站登录了，然后打开 B 网站，B 网站偷偷用你的身份向 A 网站发请求。')
r.bold = True

doc.add_paragraph('用生活中的例子理解：')
doc.add_paragraph('你去银行办业务（登录了银行系统），这时候有人递给你一张纸条（恶意网站），你接过纸条的时候，兜里的银行卡（Cookie）靠近了读卡器，自动完成了转账操作。你根本不知道发生了什么。')

p = doc.add_paragraph()
r = p.add_run('关键原因：')
r.bold = True
p.add_run('浏览器会自动带上已登录网站的 Cookie，服务器只看 Cookie 认人，不知道请求是用户主动发的还是恶意网页替用户发的。')

doc.add_page_break()

# 2. 项目中发现的CSRF漏洞
doc.add_heading('2. 项目中发现了哪些 CSRF 漏洞', level=1)

doc.add_heading('漏洞 1：注册接口无 CSRF', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🟠 高危')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x66,0x00)
doc.add_paragraph('恶意网站构造隐藏表单，向 /register 提交注册请求，在系统中创建恶意账号。')

doc.add_heading('漏洞 2：上传接口无 CSRF', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🟠 高危')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x66,0x00)
doc.add_paragraph('用户访问恶意网站时，隐藏表单自动上传恶意文件到服务器。')

doc.add_heading('漏洞 3：修改密码接口无 CSRF', level=2)
p = doc.add_paragraph()
r = p.add_run('风险等级：🔴 严重')
r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)
doc.add_paragraph('这是最危险的！恶意网站可以悄无声息地把用户密码改掉，用户下次登录就进不去了。')

doc.add_paragraph('修复前代码（以修改密码为例）：')
p = doc.add_paragraph()
r = p.add_run('''# ❌ 漏洞代码
@app.route("/change-password", methods=["POST"])
def change_password():
    username = request.form.get("username", "")
    new_password = request.form.get("new_password", "")
    # 没有任何 CSRF 校验！
    USERS[username]["password"] = bcrypt.hashpw(...)''')
r.font.size = Pt(9); r.font.name = 'Courier New'

doc.add_paragraph('修复后代码：')
p = doc.add_paragraph()
r = p.add_run('''# ✅ 安全代码
@app.route("/change-password", methods=["POST"])
def change_password():
    # 先校验 CSRF Token！
    csrf_token = request.form.get("csrf_token")
    if csrf_token != session.get("csrf_token"):
        return "CSRF token 无效", 400
    # 再执行密码修改
    ...''')
r.font.size = Pt(9); r.font.name = 'Courier New'; r.font.color.rgb = RGBColor(0x00,0x80,0x00)

doc.add_page_break()

# 3. 漏洞汇总表
doc.add_heading('3. 漏洞汇总表', level=1)
add_table(doc, ['编号','漏洞名称','风险','位置','攻击方式','后果'],
    [
        ['SEC-CSRF-01','注册接口无CSRF','🟠 高危','/register POST','隐藏表单提交','创建恶意账号'],
        ['SEC-CSRF-02','上传接口无CSRF','🟠 高危','/upload POST','隐藏表单提交','上传恶意文件'],
        ['SEC-CSRF-03','改密接口无CSRF','🔴 严重','/change-password POST','隐藏表单提交','用户密码被篡改'],
    ])

doc.add_page_break()

# 4. 修复原理
doc.add_heading('4. 修复原理：CSRF Token', level=1)
doc.add_paragraph('CSRF 防护的核心就是"暗号验证机制"。')

p = doc.add_paragraph()
r = p.add_run('CSRF Token = 只有真正的用户页面才有的随机暗号')
r.bold = True

doc.add_paragraph()
doc.add_paragraph('流程：')
doc.add_paragraph('1. 用户访问正常页面 → 服务器生成一个随机 Token，存在 session 里')
doc.add_paragraph('2. 在表单里放一个隐藏字段：<input name="csrf_token" value="随机数">')
doc.add_paragraph('3. 提交表单时，服务器检查：收到的 Token == session 里存的 Token？')
doc.add_paragraph('4. 恶意网站拿不到这个随机数 → 构造的请求被拒绝')

add_table(doc, ['对比项','修复前','修复后'],
    [
        ['校验方式','无','Token比对'],
        ['恶意网站能否伪造','✅ 可以，直接post就行','❌ 不知道Token，无法伪造'],
        ['安全性','❌ 高危','✅ 安全'],
    ])

doc.add_page_break()

# 5. 修复前后对比
doc.add_heading('5. 修复前后对比', level=1)
add_table(doc, ['接口','修复前','修复后','验证结果'],
    [
        ['/register','无CSRF校验','有Token校验','✅ 拦截无Token请求'],
        ['/upload','无CSRF校验','有Token校验','✅ 拦截无Token请求'],
        ['/change-password','无CSRF校验','有Token校验','✅ 拦截无Token请求'],
        ['/login','已有Token校验','保持不动','✅ 已有防护'],
        ['/recharge','已有Token校验','保持不动','✅ 已有防护'],
    ], hc="1A6B3C")

doc.add_page_break()

# 6. 新手总结
doc.add_heading('6. 新手总结', level=1)

tips = [
    ('每个 POST 表单都加 CSRF Token', '登录、注册、上传、修改密码、充值……只要涉及数据修改都要加'),
    ('Token 用 secrets 生成', '不要用时间戳、用户名之类可预测的值做 Token'),
    ('登录成功后刷新 Token', '防止 Token 重放攻击'),
    ('SameSite Cookie 是第二道防线', '设置 SESSION_COOKIE_SAMESITE="Lax"，限制跨站自动带 Cookie'),
]
for title, desc in tips:
    p = doc.add_paragraph()
    r = p.add_run('✅ ' + title); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph(desc + '\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('记住：没有 CSRF Token 的 POST 请求就像没锁的门，谁都能推！')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

output_path = '/home/user/projects/user-mgr/CSRF漏洞修复报告_新手版.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
